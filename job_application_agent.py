import requests
from bs4 import BeautifulSoup
from docx import Document
import pandas as pd

def job_search(keyword):
    url = f"https://www.example.com/jobs?q={keyword}"  # Replace with actual job site URL
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    
    jobs = []
    for job in soup.find_all('div', class_='job-listing'):  # Adjust according to the website's structure
        title = job.find('h2').text
        company = job.find('span', class_='company').text
        jobs.append({'title': title, 'company': company})
    
    return jobs

def edit_cv(template_path, output_path, name, job_title, skills):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        if '{name}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{name}', name)
        if '{job_title}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{job_title}', job_title)
        if '{skills}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{skills}', ', '.join(skills))
    doc.save(output_path)

def edit_cover_letter(template_path, output_path, name, job_title, company):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        if '{name}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{name}', name)
        if '{job_title}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{job_title}', job_title)
        if '{company}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{company}', company)
    doc.save(output_path)

def log_application(data):
    df = pd.DataFrame(data)
    df.to_excel('applications_log.xlsx', index=False, mode='a', header=False)

def main():
    keyword = "Software Developer"
    name = "John Doe"  # Replace with your name
    skills = ["Python", "Java", "SQL"]  # Replace with your skills

    jobs = job_search(keyword)

    for job in jobs:
        job_title = job['title']
        company = job['company']

        edit_cv('template_cv.docx', 'edited_cv.docx', name, job_title, skills)
        edit_cover_letter('template_cover_letter.docx', 'edited_cover_letter.docx', name, job_title, company)

        log_data = {
            'Name': [name],
            'Job Title': [job_title],
            'Company': [company],
            'Status': ['Applied']
        }
        log_application(log_data)

if __name__ == "__main__":
    main()